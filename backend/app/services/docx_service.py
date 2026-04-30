from __future__ import annotations

import re
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from app.models.parse_record import ParseRecord


STORAGE_DIR = Path(__file__).resolve().parents[2] / "storage"
TXT_DIR = STORAGE_DIR / "txt"
DOCX_DIR = STORAGE_DIR / "docx"


def _safe_filename(value: str) -> str:
    safe = re.sub(r"[^A-Za-z0-9_.-]+", "_", value).strip("._")
    return safe or "record"


def _record_title(record: ParseRecord) -> str:
    return record.title or f"解析记录 {record.id}"


def _record_content(record: ParseRecord) -> str:
    return record.main_content or record.clean_text or ""


def _generated_at(record: ParseRecord) -> datetime:
    return record.completed_at or datetime.utcnow()


def generate_txt(record: ParseRecord) -> str:
    TXT_DIR.mkdir(parents=True, exist_ok=True)
    file_path = TXT_DIR / f"{_safe_filename(str(record.id))}.txt"

    lines = [
        f"标题：{_record_title(record)}",
        "",
        f"来源链接：{record.final_url or record.url}",
        f"解析时间：{_generated_at(record).strftime('%Y-%m-%d %H:%M:%S')}",
        "",
        "一、网页正文",
        "",
        _record_content(record) or "未提取到正文内容。",
        "",
        "二、文件链接",
        "",
    ]

    if record.attachments:
        for index, attachment in enumerate(record.attachments, start=1):
            lines.extend(
                [
                    f"{index}. {attachment.file_name or attachment.link_text or attachment.file_url}",
                    f"   类型：{attachment.file_type or '-'}",
                    f"   链接：{attachment.file_url}",
                    "",
                ]
            )
    else:
        lines.append("未提取到文件链接。")

    file_path.write_text("\n".join(lines), encoding="utf-8")
    return str(file_path)


def generate_docx(record: ParseRecord) -> str:
    DOCX_DIR.mkdir(parents=True, exist_ok=True)
    file_path = DOCX_DIR / f"{_safe_filename(str(record.id))}.docx"

    document = Document()
    title = document.add_heading(_record_title(record), level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_paragraph(f"来源链接：{record.final_url or record.url}")
    document.add_paragraph(f"解析时间：{_generated_at(record).strftime('%Y-%m-%d %H:%M:%S')}")

    document.add_heading("一、网页正文", level=1)
    content = _record_content(record)
    if content:
        for paragraph_text in content.splitlines():
            if paragraph_text.strip():
                paragraph = document.add_paragraph(paragraph_text.strip())
                paragraph.paragraph_format.space_after = Pt(6)
    else:
        document.add_paragraph("未提取到正文内容。")

    document.add_heading("二、文件链接", level=1)
    table = document.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = table.rows[0].cells
    headers[0].text = "序号"
    headers[1].text = "文件名"
    headers[2].text = "类型"
    headers[3].text = "下载链接"

    if record.attachments:
        for index, attachment in enumerate(record.attachments, start=1):
            cells = table.add_row().cells
            cells[0].text = str(index)
            cells[1].text = attachment.file_name or attachment.link_text or "-"
            cells[2].text = attachment.file_type or "-"
            cells[3].text = attachment.file_url
    else:
        cells = table.add_row().cells
        cells[0].text = "-"
        cells[1].text = "未提取到文件链接。"
        cells[2].text = "-"
        cells[3].text = "-"

    document.save(file_path)
    return str(file_path)


def generate_record_exports(record: ParseRecord) -> tuple[str, str]:
    return generate_txt(record), generate_docx(record)
